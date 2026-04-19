"""
Scan ``bayan ul quran`` (or a custom root) for ``.txt`` exports and optionally ``.docx``,
then run the same import pipeline as ``import_bayan_segments.py`` for each file whose
surah can be inferred from the filename.

Docx bodies are extracted to ``.cache/bayan_docx_txt`` (mirroring relative paths) so
the importer always sees UTF-8 text.

Usage (repo root, venv active):

    python scripts/batch_import_bayan.py --dry-run
    python scripts/batch_import_bayan.py --root "bayan ul quran" --include-docx
"""

from __future__ import annotations

import argparse
import hashlib
import importlib.util
import sys
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
if str(REPO) not in sys.path:
    sys.path.insert(0, str(REPO))

from app.bayan_parse import surah_id_from_filename


def _load_segment_importer():
    path = Path(__file__).resolve().parent / "import_bayan_segments.py"
    spec = importlib.util.spec_from_file_location("_bayan_segment_import", path)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod.import_segments_from_file


import_segments_from_file = _load_segment_importer()


def _docx_cache_path(root: Path, docx_path: Path, cache_root: Path) -> Path:
    rel = docx_path.relative_to(root)
    safe_parent = cache_root / rel.parent
    stem = rel.stem
    if len(stem) > 120:
        h = hashlib.sha256(str(rel).encode("utf-8")).hexdigest()[:16]
        stem = f"{rel.parts[-1][:80]}_{h}"
    out = safe_parent / f"{stem}.txt"
    out.parent.mkdir(parents=True, exist_ok=True)
    return out


def extract_docx_to_txt(docx_path: Path, out_txt: Path) -> None:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "Install python-docx: pip install python-docx"
        ) from e

    doc = Document(docx_path)
    lines: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    lines.append(cell.text)
    body = "\n".join(lines)
    out_txt.write_text(body, encoding="utf-8")


_SKIP_DIR_PARTS = frozenset({".cache", "__pycache__"})


def iter_txt_files(root: Path) -> list[Path]:
    out: list[Path] = []
    for p in root.rglob("*.txt"):
        if _SKIP_DIR_PARTS.intersection(p.parts):
            continue
        if "Thumbs" in p.name:
            continue
        out.append(p)
    return sorted(out)


def iter_docx_files(root: Path) -> list[Path]:
    paths: list[Path] = []
    for p in root.rglob("*.docx"):
        if _SKIP_DIR_PARTS.intersection(p.parts):
            continue
        if p.name.startswith("~$"):
            continue
        paths.append(p)
    return sorted(paths)


def main() -> int:
    ap = argparse.ArgumentParser(description="Batch-import Bayan segments under a folder tree.")
    ap.add_argument(
        "--root",
        type=Path,
        default=REPO / "bayan ul quran",
        help="Folder to scan (default: bayan ul quran next to scripts)",
    )
    ap.add_argument("--translation-id", type=int, default=1)
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument(
        "--include-docx",
        action="store_true",
        help="Extract each .docx to .cache/bayan_docx_txt and import",
    )
    ap.add_argument(
        "--skip-unknown",
        action="store_true",
        help="Skip files whose surah cannot be inferred from filename (default: exit 1 if any)",
    )
    args = ap.parse_args()

    root = args.root.resolve()
    if not root.is_dir():
        print("Not a directory:", root, file=sys.stderr)
        return 1

    cache_root = REPO / ".cache" / "bayan_docx_txt"

    jobs: list[tuple[Path, int | None, str | None]] = []
    # (path_for_read, surah_override, filename_for_surah or None to use path.name)
    for txt in iter_txt_files(root):
        jobs.append((txt, None, None))

    if args.include_docx:
        for docx in iter_docx_files(root):
            out = _docx_cache_path(root, docx, cache_root)
            try:
                if not out.is_file() or docx.stat().st_mtime > out.stat().st_mtime:
                    extract_docx_to_txt(docx, out)
            except Exception as ex:
                print("Docx extract failed:", docx, ex, file=sys.stderr)
                if not args.skip_unknown:
                    return 1
                continue
            jobs.append((out, None, docx.name))

    hard_failures = 0
    skipped_unknown = 0

    for path, surah_ov, fname_surah in jobs:
        infer_name = fname_surah if fname_surah is not None else path.name
        if surah_ov is None and surah_id_from_filename(infer_name) is None:
            print("Skip (unknown surah from filename):", path)
            skipped_unknown += 1
            if not args.skip_unknown:
                hard_failures += 1
            continue

        code = import_segments_from_file(
            path,
            translation_id=args.translation_id,
            surah_id=surah_ov,
            dry_run=args.dry_run,
            filename_for_surah=fname_surah if fname_surah is not None else None,
        )
        if code != 0:
            hard_failures += 1

    print(
        "--- batch done: %s file(s), hard_failures=%s, skipped_unknown_surah=%s ---"
        % (len(jobs), hard_failures, skipped_unknown)
    )

    return 1 if hard_failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
